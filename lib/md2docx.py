#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generic Markdown -> styled .docx converter with a page-number footer.

Usage:
  python3 md2docx.py <input.md> [output.docx]

If output.docx is omitted, it defaults to <input-name>.docx next to the input.

Renders headings (h1-h3), bold/italic inline, tables (pipe syntax),
ordered/unordered lists, blockquotes, and horizontal rules. Adds a centered
footer with an auto-updating PAGE field (updates when opened in Word or
exported to PDF).

Requires: python3 + python-docx (`pip install python-docx`).
"""
import re
import sys
import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def die(msg):
    print(f"md2docx: {msg}", file=sys.stderr)
    sys.exit(1)


if len(sys.argv) < 2:
    die("usage: md2docx.py <input.md> [output.docx]")

SRC = os.path.abspath(sys.argv[1])
if not os.path.isfile(SRC):
    die(f"input file not found: {SRC}")
OUT = (
    os.path.abspath(sys.argv[2])
    if len(sys.argv) > 2
    else os.path.splitext(SRC)[0] + ".docx"
)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()


def add_page_number(paragraph):
    """Insert an auto-updating '第 N 页' page field into a footer paragraph."""
    run = paragraph.add_run("第 ")
    set_east_asia(run)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    set_east_asia(r)
    for el in (fld, instr, sep, t, end):
        r._r.append(el)
    run2 = paragraph.add_run(" 页")
    set_east_asia(run2)


doc = Document()
doc.add_heading(os.path.splitext(os.path.basename(SRC))[0], level=0)


# --- base styles ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_east_asia(run):
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:eastAsia"), "微软雅黑")


def add_runs_with_bold(par, text):
    """Add text to paragraph, honoring **bold** and *italic* markers."""
    for tok in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            r.italic = True
        else:
            r = par.add_run(tok)
        set_east_asia(r)


def add_body_paragraph(text, style=None):
    p = doc.add_paragraph(style=style)
    add_runs_with_bold(p, text)
    return p


def flush_table():
    global table_rows, in_table
    if not table_rows:
        in_table = False
        return
    ncols = max(len(r) for r in table_rows)
    tbl = doc.add_table(rows=len(table_rows), cols=ncols)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(table_rows):
        for ci in range(ncols):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            cp = cell.paragraphs[0]
            text = row[ci] if ci < len(row) else ""
            add_runs_with_bold(cp, text)
            if ri == 0:
                for run in cp.runs:
                    run.bold = True
    tbl.autofit = True
    table_rows = []
    in_table = False


table_rows = []
in_table = False

i = 0
while i < len(lines):
    stripped = lines[i].strip()
    if not stripped:
        i += 1
        continue
    if re.fullmatch(r"-{3,}", stripped):
        flush_table()
        doc.add_paragraph()
        i += 1
        continue
    if in_table and "|" in stripped and re.fullmatch(r"\|?[\s:|-]+\|?", stripped):
        i += 1
        continue
    if stripped.startswith("|"):
        in_table = True
        table_rows.append([c.strip() for c in stripped.strip("|").split("|")])
        i += 1
        continue
    if in_table:
        flush_table()
    if stripped.startswith("### "):
        h = doc.add_heading(level=3)
        add_runs_with_bold(h, stripped[4:])
        i += 1
        continue
    if stripped.startswith("## "):
        h = doc.add_heading(level=2)
        add_runs_with_bold(h, stripped[3:])
        i += 1
        continue
    if stripped.startswith("# "):
        h = doc.add_heading(level=1)
        add_runs_with_bold(h, stripped[2:])
        i += 1
        continue
    if stripped.startswith("> "):
        p = doc.add_paragraph(style="Intense Quote")
        add_runs_with_bold(p, stripped[2:] + " ")
        i += 1
        continue
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs_with_bold(p, m.group(2))
        i += 1
        continue
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(p, stripped[2:])
        i += 1
        continue
    p = doc.add_paragraph()
    add_runs_with_bold(p, stripped)
    i += 1

if in_table:
    flush_table()

# --- footer with page number field ---
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(fp)

doc.save(OUT)
print("Saved:", OUT)
